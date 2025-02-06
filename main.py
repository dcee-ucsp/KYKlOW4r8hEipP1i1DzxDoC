import streamlit as st
import pandas as pd
from docx import Document
from docx.shared import Pt
from io import BytesIO
import time
import random

st.title("Generador de cartas de presentación")

# Ingresar texto
st.subheader("Datos generales", divider=True)
escuela = st.radio("Escuela profesional", ["Administración de negocios", "Contabilidad"])
correlativo = st.number_input("Correlativo", step=1, min_value=0)
fecha = st.date_input("Fecha de emisión")
tipo_practicas = st.radio("Tipo de prácticas", ["Pre-profesionales", "Profesionales"])

st.subheader("Datos del alumno", divider=True)
nombre = st.text_input("Nombre del alumno")
dni_est = st.text_input("Ingrese el DNI del alumno")
genero_est = st.radio("Género del estudiante", ["Masculino", "Femenino"])
semestre_alumno = st.selectbox("Semestre", ("séptimo", "octavo", "noveno", "décimo", "egresado"))
periodo_pract = st.slider("Periodo de prácticas (meses)", 1, 12)

st.subheader("Datos del empleador", divider=True)
nombre_empresa = st.text_input("Nombre de la empresa")
referencia = st.selectbox("Referencia", ("Señor", "Señora", "Señorita", "Estimado", "Estimada"))
nombre_empleador = st.text_input("Nombre del empleador")
cargo_empleador = st.text_input("Cargo del empleador")

# Conversión de fecha
meses = {
    1: "enero", 2: "febrero", 3: "marzo", 4: "abril", 5: "mayo", 6: "junio", 
    7: "julio", 8: "agosto", 9: "septiembre", 10: "octubre", 11: "noviembre", 12: "diciembre"
}
fecha_larga = f"{fecha.day} de {meses[fecha.month]} del {fecha.year}"

# Fill correlativo
correlativo = str(int(correlativo)).zfill(3)

# Conversión género
gen_alumn = ("el alumno" if genero_est == "Masculino" else "la alumna") if escuela == "Administración de negocios" else ("El alumno" if genero_est == "Masculino" else "La alumna")

# Semestre del alumno
reemplazo_semestre = f"del {semestre_alumno} semestre" if semestre_alumno != "egresado" else ("egresado" if genero_est == "Masculino" else "egresada")

# Texto meses
meses_texto = {
    1: "un mes", 2: "dos meses", 3: "tres meses", 4: "cuatro meses", 
    5: "cinco meses", 6: "seis meses", 7: "siete meses", 8: "ocho meses", 
    9: "nueve meses", 10: "diez meses", 11: "once meses", 12: "doce meses"
}
periodo_pract_texto = meses_texto[periodo_pract]

# Identif.

identificacion = (
    "Identificado" if genero_est == "Masculino" else "Identificada"
    if escuela == "Contabilidad"
    else "identificado" if genero_est == "Masculino" else "identificada"
)

# Plantilla

if escuela == "Contabilidad":
    doc = Document("Plantillas/plantilla_cont.docx")
elif escuela == "Administración de negocios":
    doc = Document("Plantillas/plantilla_adm.docx")

# Tabla de resumen

df = pd.DataFrame([{
    "CODIGO DE CARTA": f"{'DIRADM' if escuela == 'Administración de Negocios' else 'DIRCONT'}-{correlativo}-{fecha.year}",
    "ALUMNOS": nombre,
    "A QUIEN VA DIRIGIDA": f"{referencia} {nombre_empleador}",
    "TIPO DE CARTA - ASUNTO": f"Carta de {tipo_practicas}",
    "FECHA": fecha.strftime("%d/%m/%Y")  # Formato de fecha DD/MM/YYYY
}])

def get_random_step():
    steps = [
        "Tomando un café ☕",
        "Visitando la cafetería 🍽️",
        "Lavando los platos 🍽️",
        "Regando las plantas 🌱",
        "Ordenando el escritorio 📚",
        "Alimentando al gato 🐱",
        "Haciendo ejercicio 🏃",
        "Meditando un momento 🧘",
        "Revisando el correo 📧",
        "Estirando los brazos 💪"
    ]
    return random.choice(steps)

def cook_breakfast():
    msg = st.toast('📜 Preparando el documento...')
    time.sleep(1)
    msg.toast(get_random_step())
    time.sleep(1)
    msg.toast('✅ Documento listo para descargar', icon="📄")

# Formateo

def reemplazar_texto(doc, marcador, nuevo_texto):
    for paragraph in doc.paragraphs:
        if marcador in paragraph.text:
            paragraph.text = paragraph.text.replace(marcador, str(nuevo_texto))
    
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if marcador in paragraph.text:
                        paragraph.text = paragraph.text.replace(marcador, str(nuevo_texto))

def set_font_style(doc):
    times_new_roman_font = 'Times New Roman'
    font_size = Pt(11)
    
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            run.font.name = times_new_roman_font
            run.font.size = font_size
    
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.name = times_new_roman_font
                        run.font.size = font_size

# Reemplazo de valores

dt_compt = all([correlativo, fecha, tipo_practicas, nombre, dni_est, genero_est, semestre_alumno,
    periodo_pract, nombre_empresa, referencia, nombre_empleador, cargo_empleador])

if st.button("Generar Documento", disabled=not dt_compt):
    cook_breakfast()
    
    reemplazar_texto(doc, "{{CORRELATIVO}}", correlativo)
    reemplazar_texto(doc, "{{ANIO}}", fecha.year)
    reemplazar_texto(doc, "{{FECHA_LARGA}}", fecha_larga)
    reemplazar_texto(doc, "{{REFERENCIA}}", referencia)
    reemplazar_texto(doc, "{{JEFE_DIRECTO}}", nombre_empleador)
    reemplazar_texto(doc, "{{CARGO_JEFE}}", cargo_empleador)
    reemplazar_texto(doc, "{{NOMBRE_EMPRESA}}", nombre_empresa)
    reemplazar_texto(doc, "{{GEN_ALUM}}", gen_alumn)
    reemplazar_texto(doc, "{{SEM_ALUM}}", reemplazo_semestre)
    reemplazar_texto(doc, "{{NOMBRE_ALUMNO}}", nombre)
    reemplazar_texto(doc, "{{DNI_ALUMNO}}", dni_est)
    reemplazar_texto(doc, "{{TIPO_PRACTICAS}}", tipo_practicas)
    reemplazar_texto(doc, "{{PERIODO_MESES}}", periodo_pract_texto)
    reemplazar_texto(doc, "{{IDENT}}", identificacion)
    
    set_font_style(doc)

    # Mostrar DataFrame en Streamlit
    st.dataframe(df)
    
    doc_nom = f"{'DIRADM' if escuela == 'Administración de negocios' else 'DIRCONT'} - {correlativo} - {fecha.year} - {nombre} - {nombre_empresa}"
    
    buffer_docx = BytesIO()
    doc.save(buffer_docx)
    buffer_docx.seek(0)
    
    st.download_button(
        label="📄 Descargar DOCX",
        data=buffer_docx,
        file_name=f"{doc_nom}.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
